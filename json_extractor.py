from docx import Document

# Your certifications data
certifications = [
    # JSON DATA HERE
]

# Create a Word document
doc = Document()

# Add a title
doc.add_heading("Certification Plan", level=1)

# Add a table
table = doc.add_table(rows=1, cols=len(certifications[0]))
table.style = 'Table Grid'

# Add header row
header_cells = table.rows[0].cells
for idx, header in enumerate(certifications[0].keys()):
    header_cells[idx].text = header

# Add data rows
for cert in certifications:
    row_cells = table.add_row().cells
    for idx, key in enumerate(cert):
        row_cells[idx].text = str(cert[key])

# Save the document
doc.save("Certification_Plan.docx")

print("Word document created successfully!")
