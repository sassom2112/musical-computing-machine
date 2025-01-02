from bs4 import BeautifulSoup
from docx import Document

# Load the HTML content
with open('./Security-Certification-Roadmap9.html', 'r', encoding='utf-8') as file:
    html_content = file.read()

# Parse the HTML content with BeautifulSoup
soup = BeautifulSoup(html_content, 'html.parser')

# Initialize Word document
doc = Document()
doc.add_heading('Certifications Table', level=1)

# Extract certification data
rows = []
for link in soup.find_all('a'):
    cert_name = link.get_text(strip=True)
    tooltip = link.get('tooltiptext') or link.get('tooltipright') or link.get('tooltipleft', '')
    url = link.get('href', '')
    rows.append([cert_name, tooltip, url])

# Add table to the Word document
table = doc.add_table(rows=1, cols=3)
table.style = 'Table Grid'
header_cells = table.rows[0].cells
header_cells[0].text = 'Certification Name'
header_cells[1].text = 'Description'
header_cells[2].text = 'URL'

# Populate table rows
for row_data in rows:
    row_cells = table.add_row().cells
    row_cells[0].text = row_data[0]
    row_cells[1].text = row_data[1]
    row_cells[2].text = row_data[2]

# Save the Word document
doc.save('Certifications.docx')

print("Word document created successfully!")
